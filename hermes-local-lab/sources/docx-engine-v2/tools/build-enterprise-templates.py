"""Build the two deterministic enterprise Carbone DOCX templates."""

import hashlib
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BLUE = RGBColor(31, 78, 121)
GRAY = RGBColor(89, 89, 89)
BODY_CJK_FONT = "宋体"
HEADING_CJK_FONT = "黑体"


def set_font(run, name="Arial", size=11, bold=False, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_CJK_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CJK_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover_line(doc, label, token):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Inches(1.1)
    set_font(p.add_run(f"{label}："), size=10.5, bold=True, color=GRAY)
    set_font(p.add_run(token), size=10.5)


def build(template_id, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run(subtitle), size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("第 "), size=9, color=GRAY)
    add_page_field(footer)
    set_font(footer.add_run(" 页"), size=9, color=GRAY)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(90)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    set_font(title.add_run("{d.cover.title}"), size=25, bold=True, color=BLUE)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(70)
    set_font(sub.add_run("{d.cover.subtitle}"), size=16, bold=True)
    for label, token in (
        ("客户单位", "{d.cover.client}"),
        ("签发单位", "{d.cover.issuer}"),
        ("编制单位", "{d.cover.compiler}"),
        ("文档版本", "{d.cover.version}"),
        ("密级", "{d.cover.security_level}"),
        ("编制日期", "{d.cover.date}"),
    ):
        add_cover_line(doc, label, token)

    doc.add_page_break()
    toc_title = doc.add_paragraph("目录", style="Heading 1")
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_first = doc.add_paragraph()
    set_font(toc_first.add_run("{d.sections[i].title}"), size=11)
    toc_next = doc.add_paragraph()
    set_font(toc_next.add_run("{d.sections[i+1].title}"), size=11)
    doc.add_page_break()

    doc.add_paragraph("{d.sections[i].title}", style="Heading 1")
    doc.add_paragraph("{d.sections[i].paragraphs[i].text}")
    doc.add_paragraph("{d.sections[i].paragraphs[i+1].text}")
    doc.add_paragraph("{d.sections[i+1].title}", style="Heading 1")

    core = doc.core_properties
    core.title = f"{subtitle}企业模板"
    core.subject = "approved canonical content only"
    core.author = "Taiji DOCX Engine V2"
    target = ROOT / "templates" / template_id / "template.docx"
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    package_dir = target.parent
    files = {
        name: hashlib.sha256((package_dir / name).read_bytes()).hexdigest()
        for name in ("manifest.json", "schema.json", "data-adapter.js", "template.docx")
    }
    binding = {
        "schemaVersion": "docx-template-package-binding/v1",
        "files": files,
        "packageSha256": hashlib.sha256(
            json.dumps(files, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
        ).hexdigest(),
    }
    (package_dir / "template-package.binding.json").write_text(
        json.dumps(binding, ensure_ascii=False, sort_keys=True, indent=2) + "\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    build("enterprise-work-report", "工作汇报")
    build("enterprise-research-report", "研究报告")

"""Build deterministic A4 Carbone templates for standalone expert-team delivery."""

import base64
import hashlib
import json
import re
import sys
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BLUE = RGBColor(13, 119, 145)
DARK = RGBColor(24, 57, 91)
GRAY = RGBColor(96, 112, 128)
BODY_CJK_FONT = "宋体"
HEADING_CJK_FONT = "黑体"
TRANSPARENT_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)
DETERMINISTIC_ZIP_TIMESTAMP = (1980, 1, 1, 0, 0, 0)
UNAVAILABLE_FONT_REPLACEMENTS = {
    "ＭＳ 明朝": BODY_CJK_FONT,
    "ＭＳ ゴシック": HEADING_CJK_FONT,
    "Courier": "Arial",
}
PORTABLE_FONT_FAMILIES = {"", "Arial", "Symbol", BODY_CJK_FONT, HEADING_CJK_FONT}


def set_font(run, *, cjk=BODY_CJK_FONT, western="Arial", size=11, bold=False, color=None):
    run.font.name = western
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cjk)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CJK_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35
    for style_name, size, before, after, outline_level in (
        ("Heading 1", 16, 16, 8, 0),
        ("Heading 2", 14, 12, 6, 1),
        ("Heading 3", 12, 8, 4, 2),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = DARK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        paragraph_properties = style._element.get_or_add_pPr()
        outline = paragraph_properties.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            paragraph_properties.append(outline)
        outline.set(qn("w:val"), str(outline_level))


def set_current_word_compatibility_mode(document):
    settings = document.settings._element
    compatibility = settings.find(qn("w:compat"))
    if compatibility is None:
        compatibility = OxmlElement("w:compat")
        settings.append(compatibility)
    compatibility_mode = next(
        (
            item
            for item in compatibility.findall(qn("w:compatSetting"))
            if item.get(qn("w:name")) == "compatibilityMode"
        ),
        None,
    )
    if compatibility_mode is None:
        compatibility_mode = OxmlElement("w:compatSetting")
        compatibility.append(compatibility_mode)
    compatibility_mode.set(qn("w:name"), "compatibilityMode")
    compatibility_mode.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    compatibility_mode.set(qn("w:val"), "15")


def shade_cell(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table_template(document):
    title = document.add_paragraph()
    set_font(title.add_run("{d.tables[i].title}"), cjk=HEADING_CJK_FONT, size=11, bold=True, color=DARK)
    table = document.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tokens = (
        (
            "{d.tables[i].headers.c1}",
            "{d.tables[i].headers.c2}",
            "{d.tables[i].headers.c3}",
            "{d.tables[i].headers.c4}",
        ),
        (
            "{d.tables[i].rows[i].c1}",
            "{d.tables[i].rows[i].c2}",
            "{d.tables[i].rows[i].c3}",
            "{d.tables[i].rows[i].c4}",
        ),
        (
            "{d.tables[i].rows[i+1].c1}",
            "{d.tables[i].rows[i+1].c2}",
            "{d.tables[i].rows[i+1].c3}",
            "{d.tables[i].rows[i+1].c4}",
        ),
    )
    for row_index, row_tokens in enumerate(tokens):
        for cell, token in zip(table.rows[row_index].cells, row_tokens):
            cell.width = Mm(40)
            paragraph = cell.paragraphs[0]
            set_font(paragraph.add_run(token), size=9.5, bold=row_index == 0)
            if row_index == 0:
                shade_cell(cell, "DDEBF7")
    next_title = document.add_paragraph()
    set_font(next_title.add_run("{d.tables[i+1].title}"), cjk=HEADING_CJK_FONT, size=11, bold=True, color=DARK)


def add_figure_template(document):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(BytesIO(TRANSPARENT_PNG), width=Mm(1), height=Mm(1))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(caption.add_run("{d.figures[0].title}"), size=9.5, color=GRAY)
    description = document.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(description.add_run("{d.figures[0].description}"), size=9, color=GRAY)


def write_binding(package_dir):
    files = {
        name: hashlib.sha256((package_dir / name).read_bytes()).hexdigest()
        for name in ("manifest.json", "schema.json", "data-adapter.js", "template.docx")
    }
    binding = {
        "schemaVersion": "docx-template-package-binding/v1",
        "files": files,
        "packageSha256": hashlib.sha256(
            json.dumps(files, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
        ).hexdigest(),
    }
    (package_dir / "template-package.binding.json").write_text(
        json.dumps(binding, ensure_ascii=False, sort_keys=True, indent=2) + "\n",
        encoding="utf-8",
    )


def portable_template_xml(filename, data):
    if not filename.endswith(".xml"):
        return data
    text = data.decode("utf-8")
    if filename == "word/fontTable.xml":
        text = re.sub(
            r'\s*<w:font\b(?=[^>]*w:name="([^"]+)")[^>]*(?:/>|>.*?</w:font>)',
            lambda match: match.group(0) if match.group(1) in PORTABLE_FONT_FAMILIES else "",
            text,
            flags=re.DOTALL,
        )

    def replace_font_attribute(match):
        attribute = match.group(1)
        font_name = match.group(2)
        if font_name in PORTABLE_FONT_FAMILIES:
            return match.group(0)
        replacement = UNAVAILABLE_FONT_REPLACEMENTS.get(font_name)
        if replacement is None:
            replacement = BODY_CJK_FONT if attribute == "w:eastAsia" else "Arial"
        return f'{attribute}="{replacement}"'

    text = re.sub(
        r'<w:rFonts\b[^>]*>',
        lambda tag_match: re.sub(
            r'(w:ascii|w:hAnsi|w:eastAsia|w:cs)="([^"]*)"',
            replace_font_attribute,
            tag_match.group(0),
        ),
        text,
    )
    text = re.sub(
        r'(typeface)="([^"]*)"',
        replace_font_attribute,
        text,
    )

    def chinese_language_tag(match):
        tag = match.group(0)
        for attribute in ("w:val", "w:eastAsia"):
            if re.search(rf'\b{re.escape(attribute)}="[^"]*"', tag):
                tag = re.sub(
                    rf'\b{re.escape(attribute)}="[^"]*"',
                    f'{attribute}="zh-CN"',
                    tag,
                )
            else:
                closing = "/>" if tag.endswith("/>") else ">"
                tag = tag[: -len(closing)] + f' {attribute}="zh-CN"' + closing
        return tag

    text = re.sub(r'<w:lang\b[^>]*>', chinese_language_tag, text)
    return text.encode("utf-8")


def normalize_docx_archive(docx_path):
    with ZipFile(docx_path, "r") as source:
        entries = [
            (item.filename, portable_template_xml(item.filename, source.read(item.filename)))
            for item in sorted(source.infolist(), key=lambda item: item.filename)
            if not item.is_dir()
        ]
    output = BytesIO()
    with ZipFile(output, "w", compression=ZIP_DEFLATED, compresslevel=9) as target:
        for filename, data in entries:
            info = ZipInfo(filename=filename, date_time=DETERMINISTIC_ZIP_TIMESTAMP)
            info.compress_type = ZIP_DEFLATED
            info.create_system = 3
            info.external_attr = 0o600 << 16
            target.writestr(info, data, compress_type=ZIP_DEFLATED, compresslevel=9)
    docx_path.write_bytes(output.getvalue())


def build(template_id, subtitle):
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)
    configure_styles(document)
    set_current_word_compatibility_mode(document)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run(subtitle), size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("第 "), size=9, color=GRAY)
    add_page_field(footer)
    set_font(footer.add_run(" 页"), size=9, color=GRAY)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(84)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(16)
    set_font(title.add_run("{d.document.title}"), cjk=HEADING_CJK_FONT, size=25, bold=True, color=DARK)
    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle_paragraph.add_run("{d.document.subtitle}"), cjk=HEADING_CJK_FONT, size=15, bold=True, color=BLUE)

    document.add_page_break()
    directory_title = document.add_paragraph("目录", style="Heading 1")
    directory_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    directory_first = document.add_paragraph()
    set_font(directory_first.add_run("{d.sections[i].title}"), size=11)
    directory_next = document.add_paragraph()
    set_font(directory_next.add_run("{d.sections[i+1].title}"), size=11)
    document.add_page_break()

    document.add_paragraph("{d.sections[i].title}", style="Heading 1")
    document.add_paragraph("{d.sections[i].paragraphs[i].text}")
    document.add_paragraph("{d.sections[i].paragraphs[i+1].text}")
    document.add_paragraph("{d.sections[i+1].title}", style="Heading 1")
    add_table_template(document)
    add_figure_template(document)

    core = document.core_properties
    core.title = f"{subtitle}独立版模板"
    core.subject = "standalone canonical content only"
    core.author = "Taiji DOCX Engine V2"
    package_dir = ROOT / "templates" / template_id
    package_dir.mkdir(parents=True, exist_ok=True)
    template_path = package_dir / "template.docx"
    document.save(template_path)
    normalize_docx_archive(template_path)
    write_binding(package_dir)


if __name__ == "__main__":
    templates = {
        "standalone-work-report": "工作汇报",
        "standalone-research-report": "深度研究报告",
        "standalone-meeting-minutes": "会议纪要",
        "standalone-office-material": "办公材料",
    }
    requested = sys.argv[1:] or list(templates)
    unknown = [template_id for template_id in requested if template_id not in templates]
    if unknown:
        raise SystemExit(f"unknown standalone template: {', '.join(unknown)}")
    for template_id in requested:
        build(template_id, templates[template_id])
